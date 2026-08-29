# -*- coding: utf-8 -*-
"""
Eksporter Dokumen 024 (Matriks Ekivalensi K2025 -> K2026) ke XLSX & DOCX
yang mudah dibaca oleh pembaca non-teknis.

Berbeda dari converter generik, skrip ini:
  1. Mengambil data K2025 langsung dari PDF Laporan SIAKAD (ground truth).
  2. Mengambil data K2026 dari Dokumen 005 & 007.
  3. Menerjemahkan kode kategori (E1..E5, B) menjadi kalimat tindakan yang lugas.
  4. Memberi pewarnaan konsisten per kategori dan sheet terpisah per topik.

Keluaran:
  024_RINGKAS_EKIVALENSI_UNTUK_AWAM.xlsx  (8 sheet berwarna, header beku, filter)
  024_RINGKAS_EKIVALENSI_UNTUK_AWAM.docx  (lanskap, tabel bersih, penjelasan naratif)

Jalankan: python _tools/export_024_awam.py
"""
import os
import re
import sys
from collections import defaultdict

try:
    import pdfplumber
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
except ImportError as e:
    sys.exit(f"[GAGAL] Modul belum terpasang: {e}")

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ROOT = os.path.dirname(BASE)
PDF_GT = os.path.join(ROOT, "KURIKULUM2025", "Laporan Daftar Kurikulum Prodi Sistekin.pdf")
D005 = os.path.join(BASE, "005_STRUKTUR_KURIKULUM_8_SEMESTER_DAN_PEMINATAN.md")
D007 = os.path.join(BASE, "007_FORMULASI_CPMK_DAN_SUB_CPMK_PORTFOLIO_LENGKAP.md")
D024 = os.path.join(BASE, "024_MATRIKS_EKIVALENSI_KURIKULUM2025_KE_KURIKULUM2026.md")
OUT_XLSX = os.path.join(BASE, "024_RINGKAS_EKIVALENSI_UNTUK_AWAM.xlsx")
OUT_DOCX = os.path.join(BASE, "024_RINGKAS_EKIVALENSI_UNTUK_AWAM.docx")

# ---------------- Palet warna & terjemahan kategori ----------------
NAVY = "1F3864"
HEAD_FILL = "1F3864"
KAT = {
    "E1": ("Diakui penuh", "C6EFCE", "006100",
           "Nilai mata kuliah lama dipindahkan langsung. Mahasiswa tidak perlu mengulang."),
    "E2": ("Diakui bersyarat", "FFEB9C", "9C5700",
           "Nilai diakui, tetapi mahasiswa harus mengerjakan tugas/praktikum penyetaraan lebih dahulu."),
    "E3": ("Digabung", "BDD7EE", "1F4E79",
           "Dua mata kuliah lama dilebur menjadi satu mata kuliah baru. Nilai baru = rata-rata berbobot."),
    "E4": ("Dipecah", "D9C2E9", "5B2C87",
           "Satu mata kuliah lama menurunkan dua mata kuliah baru."),
    "E5": ("Tidak ada padanan", "F8CBAD", "833C0C",
           "Mata kuliah lama dihapus. SKS tetap tercatat di transkrip sebagai kredit bebas."),
    "KG": ("Kalah prioritas", "E4DFEC", "604A7B",
           "Ada dua mata kuliah lama menuju satu mata kuliah baru, dan yang lain lebih "
           "diprioritaskan. SKS mata kuliah ini menjadi kredit bebas."),
    "B":  ("Wajib ditempuh", "FFC7CE", "9C0006",
           "Mata kuliah baru. Belum pernah ada di Kurikulum 2025, sehingga wajib diambil."),
}
THIN = Side(style="thin", color="BFBFBF")
BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)


def muat_k2025():
    """56 MK Kurikulum 2025 dari PDF Laporan SIAKAD (ground truth)."""
    gt = {}
    with pdfplumber.open(PDF_GT) as pdf:
        for page in pdf.pages:
            for line in (page.extract_text() or "").split("\n"):
                m = re.match(r"^(\d+)\.\s+([A-Z]{3}-\d{3})\s+(.+?)\s+(\d+)\s+C\s+(\d)\s+", line)
                if m:
                    gt[m.group(2)] = dict(no=int(m.group(1)), nama=m.group(3).strip(),
                                          sks=int(m.group(4)), sem=int(m.group(5)))
    return gt


def muat_k2026():
    """67 MK portofolio Kurikulum 2026 dari Dokumen 005 (paket) & 007 (elektif)."""
    d005 = open(D005, encoding="utf-8").read()
    d007 = open(D007, encoding="utf-8").read()
    k26 = {}
    for m in re.finditer(r"### SEMESTER (\d) \(\d+ SKS\)[^#]*", d005):
        s = int(m.group(1))
        for kode, nama, sks, tipe, kat in re.findall(
            r"^\| [\d.B]+ \| `([A-Z]{3}-\d{3})` \| ([^|]+?) \| (\d+) \| ([^|]+?) \| ([^|]+?) \|",
            m.group(0), re.M,
        ):
            k26[kode] = dict(nama=nama.strip(), sks=int(sks), tipe=tipe.strip(),
                             kelompok=kat.strip(), sem=s)
    for m in re.finditer(
        r"\*\*Kode & Nama Mata Kuliah\*\* \| \*\*(ST[ABC]-\d\d) — ([^*]+?)\*\*.*?"
        r"\n\| \*\*Bobot SKS / Tipe\*\* \| \*\*(\d+) SKS\*\* / Tipe: \*\*([^*]+)\*\*.*?"
        r"\n\| \*\*Semester / Rumpun MK\*\* \| \*\*Semester (\d)\*\* / (Peminatan \d[^|]*?) \|"
        r"\n\| \*\*Prasyarat Akademik\*\* \| ([^|]+?) \|", d007, re.S,
    ):
        k26[m.group(1)] = dict(nama=m.group(2).strip(), sks=int(m.group(3)),
                               tipe=m.group(4).strip(), sem=int(m.group(5)),
                               kelompok=m.group(6).strip(), prasyarat=m.group(7).strip())
    return k26


def muat_pasangan():
    """Pasangan konversi dari Bagian 8 Dokumen 024."""
    doc = open(D024, encoding="utf-8").read()
    sec8 = doc.split("## 8. MATRIKS EKIVALENSI RINGKAS")[1].split("> [!WARNING]")[0]
    return re.findall(
        r"^\| `([A-Z]{3}-\d{2,3})` \| (\d+) \| (?:`([A-Z]{3}-\d{2,3})`|—) \| (?:(\d+)|—) \| (\S+) \|",
        sec8, re.M,
    )


# ================= XLSX =================
def tulis_sheet(ws, judul, subjudul, header, baris, lebar, kolom_kat=None, wrap_cols=()):
    """Tulis satu sheet: judul, subjudul, tabel berheader beku dan berfilter."""
    ws.sheet_view.showGridLines = False
    ws["A1"] = judul
    ws["A1"].font = Font(name="Calibri", size=14, bold=True, color=NAVY)
    ws["A2"] = subjudul
    ws["A2"].font = Font(name="Calibri", size=9.5, italic=True, color="595959")
    ws["A2"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(header))
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(header))
    ws.row_dimensions[2].height = 30

    hr = 4
    for c, h in enumerate(header, 1):
        cell = ws.cell(row=hr, column=c, value=h)
        cell.font = Font(name="Calibri", size=10, bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor=HEAD_FILL)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = BORDER
    ws.row_dimensions[hr].height = 32

    for i, row in enumerate(baris):
        r = hr + 1 + i
        for c, val in enumerate(row, 1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.font = Font(name="Calibri", size=10)
            cell.border = BORDER
            cell.alignment = Alignment(
                vertical="center",
                wrap_text=(c in wrap_cols),
                horizontal="center" if isinstance(val, int) else "left",
            )
        if kolom_kat:
            label_sel = str(row[kolom_kat - 1])
            for label, bg, fg, _ in KAT.values():
                if label_sel == label:
                    cell_k = ws.cell(row=r, column=kolom_kat)
                    cell_k.fill = PatternFill("solid", fgColor=bg)
                    cell_k.font = Font(name="Calibri", size=10, bold=True, color=fg)
                    cell_k.alignment = Alignment(horizontal="center", vertical="center")
                    break

    ws.freeze_panes = ws.cell(row=hr + 1, column=1)
    if baris:
        ws.auto_filter.ref = (f"A{hr}:{get_column_letter(len(header))}{hr + len(baris)}")
    for c, w in enumerate(lebar, 1):
        ws.column_dimensions[get_column_letter(c)].width = w


# ================= PENYIAPAN DATA =================
ELEKTIF_JALUR = {"STA": "P1 Integrated Smart Systems",
                 "STB": "P2 Cloud Infrastructure & Cybersecurity",
                 "STC": "P3 Digital Platform Engineering"}
# Aturan klaim ganda berjenjang (lihat Bagian 8 Dokumen 024)
KLAIM_MENANG = {"STI-403": "STI-528", "STI-701": "STI-742"}


def siapkan_data():
    gt, k26, pasangan = muat_k2025(), muat_k2026(), muat_pasangan()

    # arah balik: MK K2026 -> daftar asal K2025
    asal = defaultdict(list)
    for lama, sks_lama, baru, sks_baru, kat in pasangan:
        if baru:
            asal[baru].append((lama, int(sks_lama), kat.split("/")[0].strip("*"), kat))

    # SHEET 2: dari sisi Kurikulum 2025 (56 MK)
    maju = []
    for kode in sorted(gt, key=lambda x: gt[x]["no"]):
        g = gt[kode]
        tuj = [(l, b, k) for l, sl, b, sb, k in pasangan if l == kode]
        if not tuj:
            continue
        for _, baru, kat in tuj:
            kk = kat.split("/")[0].strip("*")
            label = KAT[kk][0]
            if not baru:
                nama_baru, sks_baru, ket = "— (dihapus)", "", "SKS jadi kredit bebas di transkrip"
            else:
                nama_baru = k26.get(baru, {}).get("nama", "?")
                sks_baru = k26.get(baru, {}).get("sks", "")
                ket = KAT[kk][3]
                if baru in KLAIM_MENANG and kode != KLAIM_MENANG[baru]:
                    label = KAT["KG"][0]
                    ket = (f"Tidak dapat diklaim karena {KLAIM_MENANG[baru]} lebih diprioritaskan "
                           f"(capaian pembelajarannya lebih dekat). SKS menjadi kredit bebas.")
                    nama_baru = f"{nama_baru} — sudah diklaim {KLAIM_MENANG[baru]}"
            maju.append([g["no"], kode, g["nama"], g["sks"], g["sem"],
                         baru or "—", nama_baru, sks_baru or "", label, ket])

    # SHEET 3: dari sisi Kurikulum 2026 (67 MK)
    balik = []
    for kode in sorted(k26, key=lambda x: (k26[x]["sem"], x)):
        v = k26[kode]
        jalur = ELEKTIF_JALUR.get(kode[:3], "Paket wajib")
        if kode in asal:
            lst = asal[kode]
            if kode in KLAIM_MENANG:
                lst = [x for x in lst if x[0] == KLAIM_MENANG[kode]]
            kk = lst[0][2]
            label = KAT[kk][0]
            sumber = " + ".join(x[0] for x in lst)
            nama_sumber = " + ".join(gt[x[0]]["nama"] for x in lst)
            tindakan = KAT[kk][3]
        else:
            label, sumber, nama_sumber = KAT["B"][0], "—", "— (belum pernah ada)"
            tindakan = KAT["B"][3]
        balik.append([v["sem"], kode, v["nama"], v["sks"], jalur,
                      sumber, nama_sumber, label, tindakan])

    # SHEET 4: neraca per semester (paket wajib)
    punya = set(asal)
    neraca = defaultdict(lambda: {"d": 0, "f": 0, "baru": []})
    for kode, v in k26.items():
        if kode[:3] in ELEKTIF_JALUR:
            continue
        sisi = "d" if kode in punya else "f"
        neraca[v["sem"]][sisi] += v["sks"]
        if sisi == "f":
            neraca[v["sem"]]["baru"].append(f"{kode} {v['nama']}")
    baris_neraca = []
    for s in sorted(neraca):
        n = neraca[s]
        tot = n["d"] + n["f"]
        baris_neraca.append([
            f"Semester {s}", tot, n["d"], n["f"],
            f"{n['d'] / tot * 100:.1f}%" if tot else "—",
            "; ".join(n["baru"]) if n["baru"] else "— (tidak ada, semua diakui)",
        ])
    tot_d = sum(n["d"] for n in neraca.values())
    tot_f = sum(n["f"] for n in neraca.values())
    baris_neraca.append(["TOTAL PAKET WAJIB", tot_d + tot_f, tot_d, tot_f,
                         f"{tot_d / (tot_d + tot_f) * 100:.1f}%", "5 mata kuliah baru"])

    # SHEET 5: MK wajib baru
    wajib_baru = []
    for kode in sorted([c for c in k26 if c not in punya and c[:3] not in ELEKTIF_JALUR],
                       key=lambda x: k26[x]["sem"]):
        v = k26[kode]
        wajib_baru.append([v["sem"], kode, v["nama"], v["sks"], v["kelompok"]])

    # SHEET 6: elektif peminatan
    elektif = []
    for kode in sorted([c for c in k26 if c[:3] in ELEKTIF_JALUR],
                       key=lambda x: (ELEKTIF_JALUR[x[:3]], k26[x]["sem"], x)):
        v = k26[kode]
        if kode in asal:
            lst = asal[kode]
            label = KAT[lst[0][2]][0]
            sumber = " + ".join(x[0] for x in lst)
            ket = f"Diakui HANYA jika mahasiswa memilih jalur {ELEKTIF_JALUR[kode[:3]][:2]}"
        else:
            label, sumber, ket = KAT["B"][0], "—", "Mata kuliah baru, wajib ditempuh"
        elektif.append([ELEKTIF_JALUR[kode[:3]], v["sem"], kode, v["nama"], v["sks"],
                        sumber, label, ket])

    # SHEET 7: klaster peleburan
    klaster = []
    doc = open(D024, encoding="utf-8").read()
    for m in re.finditer(
        r"\| \*\*(G-\d)\*\* \| (.+?) \| (\d+) \| (.+?) \| (\d+) \| (−?\d+) \| (.+?) \|", doc
    ):
        klaster.append([m.group(1), m.group(2).replace("`", ""), int(m.group(3)),
                        m.group(4).replace("`", ""), int(m.group(5)),
                        m.group(6).replace("−", "-"), m.group(7)])

    return gt, k26, maju, balik, baris_neraca, wajib_baru, elektif, klaster, tot_d, tot_f


# ================= TULIS XLSX =================
def tulis_xlsx(data):
    gt, k26, maju, balik, neraca, wajib_baru, elektif, klaster, tot_d, tot_f = data
    wb = openpyxl.Workbook()

    # --- Sheet 1: Baca Ini Dulu ---
    ws = wb.active
    ws.title = "1. Baca Ini Dulu"
    ws.sheet_view.showGridLines = False
    ws["A1"] = "PENYETARAAN MATA KULIAH: KURIKULUM 2025 → KURIKULUM 2026"
    ws["A1"].font = Font(name="Calibri", size=16, bold=True, color=NAVY)
    ws["A2"] = "Program Studi Sistem dan Teknologi Informasi (S1) — FSTI Universitas Widyagama Malang"
    ws["A2"].font = Font(name="Calibri", size=11, italic=True, color="595959")

    penjelasan = [
        ("", ""),
        ("APA ISI BERKAS INI?", ""),
        ("", "Berkas ini menjelaskan mata kuliah Kurikulum 2025 mana yang bisa diakui "
             "(tidak perlu diulang) ketika mahasiswa berpindah ke Kurikulum 2026, dan mata "
             "kuliah mana yang wajib ditempuh karena belum pernah ada sebelumnya."),
        ("", ""),
        ("ANGKA POKOK YANG PERLU DIINGAT", ""),
        ("Kurikulum 2025", f"{len(gt)} mata kuliah, {sum(v['sks'] for v in gt.values())} SKS, "
                           "seluruhnya wajib (tidak ada mata kuliah pilihan)"),
        ("Kurikulum 2026", f"{len(k26)} mata kuliah ditawarkan ({sum(v['sks'] for v in k26.values())} SKS), "
                           "mahasiswa menempuh paket 146 SKS / 55 mata kuliah"),
        ("Yang bisa diakui", f"{tot_d} dari {tot_d + tot_f} SKS paket wajib "
                             f"({tot_d / (tot_d + tot_f) * 100:.1f}%)"),
        ("Yang wajib ditempuh", f"{tot_f} SKS pada 5 mata kuliah baru, ditambah mata kuliah "
                                "peminatan sesuai jalur yang dipilih"),
        ("Total beban tambahan", "26 SKS untuk jalur P1/P2, atau 29 SKS untuk jalur P3 "
                                 "(setara satu semester)"),
        ("", ""),
        ("ARTI ISTILAH PADA KOLOM \"STATUS\"", ""),
    ]
    r = 3
    for a, b in penjelasan:
        ws.cell(row=r, column=1, value=a).font = Font(name="Calibri", size=11, bold=bool(a and not b))
        if b:
            c = ws.cell(row=r, column=2, value=b)
            c.font = Font(name="Calibri", size=10.5)
            c.alignment = Alignment(wrap_text=True, vertical="top")
            ws.row_dimensions[r].height = 30 if len(b) > 95 else 16
        r += 1

    for label, bg, fg, arti in KAT.values():
        c1 = ws.cell(row=r, column=1, value=label)
        c1.fill = PatternFill("solid", fgColor=bg)
        c1.font = Font(name="Calibri", size=10.5, bold=True, color=fg)
        c1.alignment = Alignment(horizontal="center", vertical="center")
        c1.border = BORDER
        c2 = ws.cell(row=r, column=2, value=arti)
        c2.font = Font(name="Calibri", size=10.5)
        c2.alignment = Alignment(wrap_text=True, vertical="center")
        ws.row_dimensions[r].height = 28
        r += 1

    r += 1
    ws.cell(row=r, column=1, value="ISI TIAP LEMBAR (SHEET)").font = Font(
        name="Calibri", size=11, bold=True)
    r += 1
    daftar_sheet = [
        ("2. Dari Kurikulum 2025", "Buka bila ingin tahu: \"mata kuliah lama saya dikonversi ke mana?\""),
        ("3. Ke Kurikulum 2026", "Buka bila ingin tahu: \"mata kuliah baru ini diakui dari mata kuliah lama apa?\""),
        ("4. Neraca per Semester", "Ringkasan berapa SKS diakui dan berapa yang wajib ditempuh, per semester"),
        ("5. Wajib Ditempuh", "Daftar 5 mata kuliah baru yang tidak bisa diakui dari kurikulum lama"),
        ("6. Mata Kuliah Peminatan", "18 mata kuliah peminatan dan syarat pengakuannya per jalur"),
        ("7. Mata Kuliah Digabung", "4 pasang mata kuliah lama yang dilebur menjadi satu"),
        ("8. Sumber Data", "Dari mana angka-angka ini diambil dan bagaimana diverifikasi"),
    ]
    for nama, isi in daftar_sheet:
        ws.cell(row=r, column=1, value=nama).font = Font(name="Calibri", size=10.5, bold=True, color=NAVY)
        c = ws.cell(row=r, column=2, value=isi)
        c.font = Font(name="Calibri", size=10.5)
        c.alignment = Alignment(wrap_text=True, vertical="center")
        r += 1

    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 108

    # --- Sheet 2 ---
    tulis_sheet(
        wb.create_sheet("2. Dari Kurikulum 2025"),
        "DARI KURIKULUM 2025: MATA KULIAH LAMA SAYA DIKONVERSI KE MANA?",
        "Urut sesuai nomor pada Laporan SIAKAD Kurikulum 2025. Gunakan tombol filter pada baris header "
        "untuk menyaring, misalnya menampilkan hanya yang berstatus \"Diakui bersyarat\".",
        ["No", "Kode Lama", "Nama Mata Kuliah Kurikulum 2025", "SKS", "Smt",
         "Kode Baru", "Nama Mata Kuliah Kurikulum 2026", "SKS", "Status", "Arti bagi Mahasiswa"],
        maju, [5, 11, 42, 5, 5, 11, 42, 5, 18, 60], kolom_kat=9, wrap_cols=(3, 7, 10),
    )

    # --- Sheet 3 ---
    tulis_sheet(
        wb.create_sheet("3. Ke Kurikulum 2026"),
        "KE KURIKULUM 2026: MATA KULIAH BARU INI DIAKUI DARI MATA KULIAH LAMA APA?",
        "Urut per semester Kurikulum 2026. Lembar ini dipakai Dosen Penasihat Akademik saat menyusun "
        "Kartu Rencana Studi mahasiswa pindahan kurikulum.",
        ["Smt", "Kode", "Nama Mata Kuliah Kurikulum 2026", "SKS", "Kelompok",
         "Asal Kode Lama", "Asal Nama Mata Kuliah Kurikulum 2025", "Status", "Tindakan yang Diperlukan"],
        balik, [5, 11, 44, 5, 30, 16, 44, 18, 58], kolom_kat=8, wrap_cols=(3, 7, 9),
    )

    # --- Sheet 4 ---
    tulis_sheet(
        wb.create_sheet("4. Neraca per Semester"),
        "NERACA PER SEMESTER: BERAPA YANG DIAKUI, BERAPA YANG WAJIB DITEMPUH",
        "Hanya mata kuliah paket wajib (belum termasuk mata kuliah peminatan). Semester 4, 5, dan 8 "
        "terekognisi seluruhnya sehingga mahasiswa tidak perlu menempuh apa pun di sana.",
        ["Semester", "SKS Paket", "Diakui", "Wajib Ditempuh", "Persen Diakui",
         "Mata Kuliah yang Wajib Ditempuh"],
        neraca, [20, 11, 9, 15, 13, 62], wrap_cols=(6,),
    )

    # --- Sheet 5 ---
    tulis_sheet(
        wb.create_sheet("5. Wajib Ditempuh"),
        "LIMA MATA KULIAH YANG WAJIB DITEMPUH SEMUA MAHASISWA TRANSISI",
        "Kelima mata kuliah ini belum pernah ada pada Kurikulum 2025, sehingga tidak dapat diakui "
        "dari nilai lama. Dua di antaranya (STI-103 dan STI-307) menjadi prasyarat mata kuliah lain, "
        "sehingga harus disisipkan lebih awal.",
        ["Smt", "Kode", "Nama Mata Kuliah", "SKS", "Kelompok"],
        wajib_baru, [6, 11, 54, 6, 16], wrap_cols=(3,),
    )

    # --- Sheet 6 ---
    tulis_sheet(
        wb.create_sheet("6. Mata Kuliah Peminatan"),
        "MATA KULIAH PEMINATAN: 3 JALUR, MASING-MASING 6 MATA KULIAH (18 SKS)",
        "Mahasiswa memilih satu jalur dan menempuh seluruh 6 mata kuliahnya: 1 di Semester 5, "
        "2 di Semester 6, dan 3 di Semester 7. Pengakuan nilai lama hanya berlaku bila mata kuliah "
        "peminatannya sesuai jalur yang dipilih.",
        ["Jalur Peminatan", "Smt", "Kode", "Nama Mata Kuliah", "SKS",
         "Asal Kode Lama", "Status", "Syarat Pengakuan"],
        elektif, [34, 5, 10, 46, 5, 16, 18, 46], kolom_kat=7, wrap_cols=(4, 8),
    )

    # --- Sheet 7 ---
    tulis_sheet(
        wb.create_sheet("7. Mata Kuliah Digabung"),
        "EMPAT PASANG MATA KULIAH LAMA YANG DILEBUR MENJADI SATU",
        "Peleburan dilakukan untuk menghapus materi yang berulang. Nilai mata kuliah baru dihitung "
        "sebagai rata-rata berbobot SKS dari nilai kedua mata kuliah lama. Selisih SKS "
        "(total 7 SKS) dicatat sebagai kredit bebas.",
        ["Klaster", "Mata Kuliah Kurikulum 2025 yang Dilebur", "SKS Lama",
         "Menjadi Mata Kuliah Kurikulum 2026", "SKS Baru", "Selisih", "Alasan Peleburan"],
        klaster, [9, 58, 10, 42, 10, 9, 66], wrap_cols=(2, 4, 7),
    )

    # --- Sheet 8 ---
    ws8 = wb.create_sheet("8. Sumber Data")
    ws8.sheet_view.showGridLines = False
    ws8["A1"] = "SUMBER DATA DAN CARA VERIFIKASI"
    ws8["A1"].font = Font(name="Calibri", size=14, bold=True, color=NAVY)
    sumber = [
        ("Data Kurikulum 2025", "KURIKULUM2025/Laporan Daftar Kurikulum Prodi Sistekin.pdf — "
         "laporan resmi SIAKAD, dicetak 05 Agustus 2026. Inilah satu-satunya sumber sah "
         "untuk kode, nama, SKS, dan semester mata kuliah lama."),
        ("Data Kurikulum 2026", "Dokumen 005 (Struktur 8 Semester) untuk mata kuliah paket wajib, "
         "dan Dokumen 007 (Formulasi CPMK) untuk mata kuliah peminatan."),
        ("Pemetaan penyetaraan", "Dokumen 024 Bagian 3, 3A, dan 8."),
        ("Cara verifikasi", "Skrip _tools/verify_k2025_ground_truth.py membaca PDF SIAKAD secara "
         "langsung, lalu memeriksa 17 butir kecocokan. Skrip gagal bila ada satu saja penyimpangan."),
        ("Hasil verifikasi terakhir", "LULUS seluruh 17 butir: 56 mata kuliah / 146 SKS cocok, "
         "tidak ada mata kuliah lama yang terlewat, dan tidak ada kode fiktif."),
        ("Catatan penting", "Notulensi rapat BUKAN sumber data mata kuliah. Bila terjadi perbedaan "
         "angka, PDF Laporan SIAKAD yang dimenangkan."),
    ]
    r = 3
    for a, b in sumber:
        ws8.cell(row=r, column=1, value=a).font = Font(name="Calibri", size=10.5, bold=True, color=NAVY)
        c = ws8.cell(row=r, column=2, value=b)
        c.font = Font(name="Calibri", size=10.5)
        c.alignment = Alignment(wrap_text=True, vertical="top")
        ws8.row_dimensions[r].height = 42
        r += 1
    ws8.column_dimensions["A"].width = 26
    ws8.column_dimensions["B"].width = 104

    wb.save(OUT_XLSX)
    print(f"  -> {os.path.basename(OUT_XLSX)} ({len(wb.sheetnames)} sheet)")


# ================= TULIS DOCX =================
def shade(cell, hexfill):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hexfill}"/>'))


def sel(cell, teks, size=8.5, bold=False, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(teks) if teks not in (None, "") else "—")
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)


def tabel_docx(doc, header, baris, lebar, kolom_kat=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        sel(c, h, size=size, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF),
            align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(c, HEAD_FILL)
    t.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(size)
    for row in baris:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            is_kat = kolom_kat and (i == kolom_kat - 1)
            warna = None
            tebal = False
            if is_kat:
                for label, bg, fg, _ in KAT.values():
                    if str(val) == label:
                        shade(cells[i], bg)
                        warna = RGBColor(int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
                        tebal = True
                        break
            sel(cells[i], val, size=size, bold=tebal, color=warna,
                align=WD_ALIGN_PARAGRAPH.CENTER if isinstance(val, int) or is_kat else None)
    for i, w in enumerate(lebar):
        for row in t.rows:
            row.cells[i].width = Cm(w)
    return t


def judul_bab(doc, teks, size=13):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(teks)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def paragraf(doc, teks, size=10, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(teks)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.italic = italic
    r.font.color.rgb = RGBColor(0x26, 0x26, 0x26)


def tulis_docx(data):
    gt, k26, maju, balik, neraca, wajib_baru, elektif, klaster, tot_d, tot_f = data
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(1.4))

    # Halaman judul
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PENYETARAAN MATA KULIAH\nKURIKULUM 2025 → KURIKULUM 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Program Studi Sistem dan Teknologi Informasi (S1)\n"
                  "Fakultas Sains dan Teknologi Informasi — Universitas Widyagama Malang")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    judul_bab(doc, "Apa isi dokumen ini?")
    paragraf(doc, "Dokumen ini menjelaskan mata kuliah Kurikulum 2025 mana yang dapat diakui "
                  "(sehingga mahasiswa tidak perlu mengulang) ketika berpindah ke Kurikulum 2026, "
                  "serta mata kuliah mana yang wajib ditempuh karena belum pernah ada sebelumnya. "
                  "Seluruh data Kurikulum 2025 dikutip langsung dari Laporan Daftar Kurikulum Prodi "
                  "yang dicetak dari SIAKAD, sehingga dapat dipertanggungjawabkan.")

    judul_bab(doc, "Angka pokok")
    tabel_docx(doc,
               ["Pokok Bahasan", "Keterangan"],
               [["Kurikulum 2025", f"{len(gt)} mata kuliah, {sum(v['sks'] for v in gt.values())} SKS, "
                                   "seluruhnya wajib (tidak ada mata kuliah pilihan)"],
                ["Kurikulum 2026", f"{len(k26)} mata kuliah ditawarkan "
                                   f"({sum(v['sks'] for v in k26.values())} SKS); mahasiswa menempuh "
                                   "paket 146 SKS / 55 mata kuliah"],
                ["Yang dapat diakui", f"{tot_d} dari {tot_d + tot_f} SKS paket wajib "
                                      f"({tot_d / (tot_d + tot_f) * 100:.1f} persen)"],
                ["Yang wajib ditempuh", f"{tot_f} SKS pada 5 mata kuliah baru, ditambah mata kuliah "
                                        "peminatan sesuai jalur yang dipilih"],
                ["Total beban tambahan", "26 SKS bagi jalur P1 atau P2, dan 29 SKS bagi jalur P3 — "
                                         "setara satu semester"]],
               [5.5, 21.0], size=10)

    judul_bab(doc, "Arti istilah pada kolom Status")
    tabel_docx(doc, ["Status", "Artinya bagi Mahasiswa"],
               [[label, arti] for label, _, _, arti in KAT.values()],
               [4.5, 22.0], kolom_kat=1, size=10)

    doc.add_page_break()
    judul_bab(doc, "Tabel 1. Neraca per semester: berapa yang diakui, berapa yang wajib ditempuh")
    paragraf(doc, "Tabel ini mencakup mata kuliah paket wajib saja, belum termasuk mata kuliah "
                  "peminatan. Semester 4, 5, dan 8 terekognisi seluruhnya, sehingga mahasiswa "
                  "tidak perlu menempuh mata kuliah apa pun pada semester tersebut.", size=9.5)
    tabel_docx(doc,
               ["Semester", "SKS Paket", "Diakui", "Wajib Ditempuh", "Persen Diakui",
                "Mata Kuliah yang Wajib Ditempuh"],
               neraca, [2.6, 2.2, 1.8, 2.6, 2.4, 14.8], size=9)

    judul_bab(doc, "Tabel 2. Lima mata kuliah yang wajib ditempuh semua mahasiswa transisi")
    paragraf(doc, "Kelima mata kuliah berikut belum pernah ada pada Kurikulum 2025 sehingga tidak "
                  "dapat diakui dari nilai lama. Dua di antaranya, yaitu STI-103 Arsitektur dan "
                  "Organisasi Sistem Teknologi Informasi serta STI-307 Jaringan Komputer, menjadi "
                  "prasyarat bagi mata kuliah lain. Keduanya harus disisipkan paling lambat pada "
                  "Semester 5 agar tidak menghambat mata kuliah lanjutan.", size=9.5)
    tabel_docx(doc, ["Smt", "Kode", "Nama Mata Kuliah", "SKS", "Kelompok"],
               wajib_baru, [1.6, 3.0, 14.0, 1.6, 6.2], size=9.5)

    judul_bab(doc, "Tabel 3. Empat pasang mata kuliah lama yang dilebur menjadi satu")
    paragraf(doc, "Peleburan dilakukan untuk menghapus materi yang berulang. Nilai mata kuliah baru "
                  "dihitung sebagai rata-rata berbobot SKS dari nilai kedua mata kuliah lama, dan "
                  "selisih SKS sebanyak 7 SKS dicatat sebagai kredit bebas pada transkrip.", size=9.5)
    tabel_docx(doc, ["Klaster", "Mata Kuliah Kurikulum 2025 yang Dilebur", "SKS Lama",
                     "Menjadi Mata Kuliah Kurikulum 2026", "SKS Baru", "Selisih", "Alasan"],
               klaster, [1.7, 7.6, 1.6, 5.6, 1.6, 1.5, 6.8], size=8)

    doc.add_page_break()
    judul_bab(doc, "Tabel 4. Dari Kurikulum 2025: mata kuliah lama saya dikonversi ke mana?")
    paragraf(doc, "Urut sesuai nomor pada Laporan SIAKAD Kurikulum 2025. Baris berstatus "
                  "\"Tidak ada padanan\" berarti SKS tetap tercatat pada transkrip sebagai kredit "
                  "bebas, tetapi tidak mengurangi kewajiban paket 146 SKS.", size=9.5)
    tabel_docx(doc,
               ["No", "Kode Lama", "Nama Mata Kuliah Kurikulum 2025", "SKS", "Smt",
                "Kode Baru", "Nama Mata Kuliah Kurikulum 2026", "SKS", "Status"],
               [r[:9] for r in maju], [1.1, 2.2, 6.6, 1.1, 1.1, 2.2, 6.6, 1.1, 3.4],
               kolom_kat=9, size=7.5)

    doc.add_page_break()
    judul_bab(doc, "Tabel 5. Ke Kurikulum 2026: mata kuliah baru ini diakui dari mata kuliah lama apa?")
    paragraf(doc, "Urut per semester Kurikulum 2026. Tabel ini dipakai Dosen Penasihat Akademik "
                  "saat menyusun Kartu Rencana Studi mahasiswa yang berpindah kurikulum.", size=9.5)
    tabel_docx(doc,
               ["Smt", "Kode", "Nama Mata Kuliah Kurikulum 2026", "SKS", "Kelompok",
                "Asal Kode Lama", "Status"],
               [[r[0], r[1], r[2], r[3], r[4], r[5], r[7]] for r in balik],
               [1.1, 2.2, 7.4, 1.1, 6.4, 3.4, 3.4], kolom_kat=7, size=7.5)

    doc.add_page_break()
    judul_bab(doc, "Tabel 6. Mata kuliah peminatan: tiga jalur, masing-masing enam mata kuliah")
    paragraf(doc, "Mahasiswa memilih satu jalur peminatan dan menempuh seluruh enam mata kuliahnya: "
                  "satu pada Semester 5, dua pada Semester 6, dan tiga pada Semester 7. Pengakuan "
                  "nilai lama hanya berlaku apabila mata kuliah peminatan tersebut berada pada jalur "
                  "yang dipilih mahasiswa; bila tidak, SKS-nya menjadi kredit bebas.", size=9.5)
    tabel_docx(doc,
               ["Jalur Peminatan", "Smt", "Kode", "Nama Mata Kuliah", "SKS",
                "Asal Kode Lama", "Status"],
               [[r[0], r[1], r[2], r[3], r[4], r[5], r[6]] for r in elektif],
               [6.4, 1.1, 2.2, 7.4, 1.1, 3.2, 3.6], kolom_kat=7, size=7.5)

    judul_bab(doc, "Sumber data dan cara verifikasi")
    tabel_docx(doc, ["Butir", "Keterangan"],
               [["Data Kurikulum 2025", "Laporan Daftar Kurikulum Prodi Sistekin (SIAKAD), dicetak "
                                        "05 Agustus 2026. Inilah satu-satunya sumber sah untuk kode, "
                                        "nama, SKS, dan semester mata kuliah lama."],
                ["Data Kurikulum 2026", "Dokumen 005 untuk mata kuliah paket wajib, Dokumen 007 "
                                        "untuk mata kuliah peminatan."],
                ["Pemetaan penyetaraan", "Dokumen 024 Bagian 3, 3A, dan 8."],
                ["Cara verifikasi", "Skrip verify_k2025_ground_truth.py membaca PDF SIAKAD secara "
                                    "langsung dan memeriksa 17 butir kecocokan; proses gagal bila "
                                    "ada satu saja penyimpangan."],
                ["Hasil terakhir", "LULUS seluruh 17 butir. 56 mata kuliah / 146 SKS cocok, tidak "
                                   "ada mata kuliah lama yang terlewat, tidak ada kode fiktif."],
                ["Catatan", "Notulensi rapat bukan sumber data mata kuliah. Bila terdapat perbedaan "
                            "angka, Laporan SIAKAD yang dimenangkan."]],
               [5.0, 21.5], size=9.5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Tim Pengembang Kurikulum FSTI Universitas Widyagama Malang\n"
                  "Ringkasan dari Dokumen Resmi 024 — Kurikulum OBE Revisi SISTEKIN 2026")
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.save(OUT_DOCX)
    print(f"  -> {os.path.basename(OUT_DOCX)}")


def main():
    print("Menyusun ringkasan Dokumen 024 untuk pembaca non-teknis...")
    data = siapkan_data()
    gt, k26, maju, balik, neraca, wajib_baru, elektif, klaster, tot_d, tot_f = data
    print(f"  Kurikulum 2025 : {len(gt)} MK / {sum(v['sks'] for v in gt.values())} SKS (PDF SIAKAD)")
    print(f"  Kurikulum 2026 : {len(k26)} MK / {sum(v['sks'] for v in k26.values())} SKS")
    print(f"  Baris konversi : {len(maju)} (arah maju), {len(balik)} (arah balik)")
    print(f"  Neraca paket   : {tot_d} SKS diakui, {tot_f} SKS wajib ditempuh")
    tulis_xlsx(data)
    tulis_docx(data)
    print("Selesai.")


if __name__ == "__main__":
    main()

