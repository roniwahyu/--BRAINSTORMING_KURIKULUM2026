# -*- coding: utf-8 -*-
"""
Markdown to DOCX Converter Engine
Program Studi Sistem dan Teknologi Informasi (SISTEKIN) FSTI UWG
Mengonversi dokumen Markdown kurikulum menjadi dokumen Microsoft Word (.docx) berformat rapi, formal & siap cetak.
"""

import os
import re
import glob
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

WORKDIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# Palet Warna Formal FSTI UWG
COLOR_PRIMARY_NAVY = RGBColor(31, 56, 100)      # #1F3864
COLOR_SECONDARY_BLUE = RGBColor(46, 117, 182)   # #2E75B6
COLOR_DARK_TEXT = RGBColor(38, 38, 38)          # #262626
COLOR_MUTED_GRAY = RGBColor(89, 89, 89)         # #595959

HEX_PRIMARY_NAVY = "1F3864"
HEX_LIGHT_BLUE_BG = "F2F5F9"
HEX_BORDER_GRAY = "D3D3D3"
HEX_CALLOUT_BG = "EEF4FB"

def set_cell_background(cell, fill_hex):
    """Menyetel warna latar belakang cell tabel."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Menyetel margin/padding dalam cell tabel (dalam dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3", sz="4", val="single"):
    """Menyetel border tabel tipis dan rapi."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_callout_box(doc, text):
    """Membuat callout box berarsir untuk blockquote / alert."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, HEX_CALLOUT_BG)
    set_cell_margins(cell, top=120, bottom=120, left=200, right=150)
    
    # Border hanya di sebelah kiri (tebal)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="2E75B6"/>'
        f'<w:top w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = COLOR_PRIMARY_NAVY
    
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def format_inline_runs(paragraph, text, base_font_size=10.5, is_italic=False):
    """Memproses format inline markdown (*bold*, _italic_, `code`, math) ke runs docx."""
    pattern = r'(\*\*.*?\*\*|\*.*?\*|`.*?`|\$.*?\$)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if not token:
            continue
            
        run = paragraph.add_run()
        run.font.name = 'Calibri'
        run.font.size = Pt(base_font_size)
        run.font.color.rgb = COLOR_DARK_TEXT
        if is_italic:
            run.font.italic = True
            
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run.text = token[2:-2]
            run.font.bold = True
        elif token.startswith('*') and token.endswith('*') and len(token) >= 2:
            run.text = token[1:-1]
            run.font.italic = True
        elif token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run.text = token[1:-1]
            run.font.name = 'Consolas'
            run.font.size = Pt(base_font_size - 0.5)
            run.font.color.rgb = RGBColor(160, 40, 40)
        elif token.startswith('$') and token.endswith('$') and len(token) >= 2:
            clean_math = token[1:-1].replace(r'\text', '').replace('{', '').replace('}', '').replace(r'\ge', '≥').replace(r'\le', '≤').replace(r'\times', '×').replace(r'\rightarrow', '→')
            run.text = clean_math
            run.font.italic = True
        else:
            run.text = token

def parse_markdown_to_docx(md_path, docx_path):
    """Fungsi utama pengonversi satu file Markdown ke format DOCX profesional."""
    print(f"  -> Mengonversi: {os.path.basename(md_path)} -> {os.path.basename(docx_path)}")
    
    with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
        md_text = f.read()

    doc = Document()
    
    # Konfigurasi Halaman (A4 Portrait, Margin Normal 1 Inch)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    
    # Normal Style Configuration
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = COLOR_DARK_TEXT
    
    lines = md_text.splitlines()
    in_table = False
    table_lines = []
    in_code_block = False
    code_lines = []
    
    def flush_table():
        nonlocal in_table, table_lines
        if not table_lines:
            in_table = False
            return
            
        rows_data = []
        for t_line in table_lines:
            if re.match(r'^\s*\|[-| :]+\|\s*$', t_line):
                continue
            cells = [c.strip() for c in t_line.strip('|').split('|')]
            rows_data.append(cells)
            
        if rows_data:
            n_rows = len(rows_data)
            n_cols = max(len(r) for r in rows_data)
            
            table = doc.add_table(rows=n_rows, cols=n_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True
            set_table_borders(table, color=HEX_BORDER_GRAY, sz="4")
            
            # Header Row Repeat (CantSplit & TblHeader)
            trPr = table.rows[0]._tr.get_or_add_trPr()
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            
            for r_idx, row_cells in enumerate(rows_data):
                row = table.rows[r_idx]
                is_header = (r_idx == 0)
                
                # CantSplit per row
                r_trPr = row._tr.get_or_add_trPr()
                r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                
                for c_idx in range(n_cols):
                    cell = row.cells[c_idx]
                    cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    if is_header:
                        set_cell_background(cell, HEX_PRIMARY_NAVY)
                        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.05
                        
                        run = p.add_run(cell_text.replace('**', '').replace('`', ''))
                        run.font.name = 'Calibri'
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        if r_idx % 2 == 1:
                            set_cell_background(cell, "FFFFFF")
                        else:
                            set_cell_background(cell, HEX_LIGHT_BLUE_BG)
                            
                        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.line_spacing = 1.05
                        format_inline_runs(p, cell_text, base_font_size=9.5)
                        
            # Spacing after table
            p_after = doc.add_paragraph()
            p_after.paragraph_format.space_after = Pt(6)
            
        table_lines = []
        in_table = False

    def flush_code():
        nonlocal in_code_block, code_lines
        if not code_lines:
            in_code_block = False
            return
            
        text = "\n".join(code_lines)
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F8F9FA")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        set_table_borders(tbl, color="E2E8F0", sz="4")
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(40, 40, 40)
        
        p_after = doc.add_paragraph()
        p_after.paragraph_format.space_after = Pt(4)
        
        code_lines = []
        in_code_block = False

    for line in lines:
        stripped = line.strip()
        
        # Check code blocks
        if stripped.startswith('```'):
            if in_code_block:
                flush_code()
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue
            
        # Check tables
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(stripped)
            continue
        else:
            if in_table:
                flush_table()
                
        # Empty line
        if not stripped:
            continue
            
        # Horizontal rule
        if stripped in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("_________________________________________________________________________________")
            run.font.color.rgb = RGBColor(200, 210, 225)
            run.font.size = Pt(8)
            continue
            
        # Headings
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[2:])
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            continue
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[3:])
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            continue
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(9)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[4:])
            run.font.name = 'Calibri'
            run.font.size = Pt(11.5)
            run.font.bold = True
            run.font.color.rgb = COLOR_SECONDARY_BLUE
            continue
        elif stripped.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped[5:])
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = COLOR_PRIMARY_NAVY
            continue
            
        # Blockquote / Alert
        if stripped.startswith('>'):
            clean_quote = re.sub(r'^>\s*(\[!.*?\])?\s*', '', stripped)
            make_callout_box(doc, clean_quote)
            continue
            
        # Bullet list
        if stripped.startswith(('* ', '- ', '+ ')):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_inline_runs(p, stripped[2:])
            continue
            
        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_inline_runs(p, num_match.group(2))
            continue
            
        # Regular Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        format_inline_runs(p, stripped)
        
    if in_table:
        flush_table()
    if in_code_block:
        flush_code()
        
    doc.save(docx_path)
    print(f"     [SUKSES] Tersimpan: {os.path.getsize(docx_path)} bytes")

def convert_all():
    print("=====================================================================")
    print("  MARKDOWN TO DOCX CONVERTER ENGINE — SISTEKIN 2026")
    print("=====================================================================")
    
    md_files = sorted(glob.glob(os.path.join(WORKDIR, "*.md")))
    success_count = 0
    
    for md_path in md_files:
        base_name = os.path.splitext(os.path.basename(md_path))[0]
        docx_path = os.path.join(WORKDIR, f"{base_name}.docx")
        try:
            parse_markdown_to_docx(md_path, docx_path)
            success_count += 1
        except Exception as e:
            print(f"     [ERROR] Gagal mengonversi {os.path.basename(md_path)}: {e}")
            
    print(f"\n[SELESAI] Berhasil mengonversi {success_count}/{len(md_files)} file Markdown ke DOCX!")

if __name__ == '__main__':
    convert_all()
